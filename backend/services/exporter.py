import os
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors


def _srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def export_srt(segments: list, title: str = "") -> bytes:
    lines = []
    for i, seg in enumerate(segments, 1):
        lines.append(str(i))
        lines.append(f"{_srt_time(seg['start'])} --> {_srt_time(seg['end'])}")
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_vtt(segments: list, title: str = "") -> bytes:
    lines = ["WEBVTT", ""]
    for seg in segments:
        lines.append(f"{_vtt_time(seg['start'])} --> {_vtt_time(seg['end'])}")
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_txt(segments: list, title: str = "") -> bytes:
    lines = []
    if title:
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")
    for seg in segments:
        lines.append(f"[{seg['start_formatted']} --> {seg['end_formatted']}]")
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_docx(segments: list, title: str = "") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for seg in segments:
        ts_para = doc.add_paragraph()
        ts_run = ts_para.add_run(f"{seg['start_formatted']} → {seg['end_formatted']}")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
        ts_run.bold = True
        ts_para.paragraph_format.space_after = Pt(2)

        text_para = doc.add_paragraph(seg["text"])
        text_para.paragraph_format.space_after = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_pdf(segments: list, title: str = "") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=20,
    )
    timestamp_style = ParagraphStyle(
        "Timestamp",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#888899"),
        spaceBefore=6,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )
    text_style = ParagraphStyle(
        "SegText",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=10,
        leading=16,
    )

    story = []
    if title:
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * cm))

    for seg in segments:
        story.append(
            Paragraph(f"{seg['start_formatted']} → {seg['end_formatted']}", timestamp_style)
        )
        story.append(Paragraph(seg["text"], text_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()
